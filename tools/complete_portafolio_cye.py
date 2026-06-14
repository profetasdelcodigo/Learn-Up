from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SRC = Path(r"E:\EPT\Portafolio Cat. A - CyE 2026.docx")
OUT = Path(r"E:\EPT\Portafolio Cat. A - CyE 2026 - COMPLETADO.docx")

IMG_ROOT = Path(r"C:\Users\profe\Learn Up")
EVIDENCE_ROOT = Path(r"C:\Users\profe\Pictures\Proyecto Learn Up")
EPT_ROOT = Path(r"E:\EPT")


def set_cell(cell, text: str, *, size: int = 8, bold: bool = False) -> None:
    text = text.strip()
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def append_text(cell, text: str, *, size: int = 8, bold: bool = False) -> None:
    p = cell.add_paragraph()
    r = p.add_run(text.strip())
    r.bold = bold
    r.font.size = Pt(size)


def add_image(cell, path: Path, *, width: float = 2.45, caption: str | None = None) -> None:
    if not path.exists():
        set_cell(cell, f"PENDIENTE: insertar evidencia. No se encontró {path.name}.", size=8)
        return
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    if caption:
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.font.size = Pt(7)


def fill_row(table, row_idx: int, values: list[str], *, size: int = 8) -> None:
    row = table.rows[row_idx]
    for idx, value in enumerate(values):
        if idx < len(row.cells):
            set_cell(row.cells[idx], value, size=size)


def duplicate_row(table, source_idx: int) -> None:
    tr = table.rows[source_idx]._tr
    new_tr = deepcopy(tr)
    tr.addnext(new_tr)


def fill_evidence_table(table, rows: list[tuple[Path | None, str]]) -> None:
    # Preserve the table design: use existing rows and only add content.
    for i, (img, desc) in enumerate(rows, start=1):
        if i >= len(table.rows):
            duplicate_row(table, len(table.rows) - 1)
        cells = table.rows[i].cells
        if img:
            add_image(cells[0], img, width=2.35)
        else:
            set_cell(cells[0], "PENDIENTE", size=8, bold=True)
        if len(cells) > 1:
            set_cell(cells[1], desc, size=8)


def polish_spanish_text(doc: Document) -> None:
    replacements = {
        "San Jose": "San José",
        "analisis": "análisis",
        "autenticacion": "autenticación",
        "colaboracion": "colaboración",
        "organizacion": "organización",
        "modulos": "módulos",
        "examenes": "exámenes",
        "academico": "académico",
        "academica": "académica",
        "acompanamiento": "acompañamiento",
        "companeros": "compañeros",
        "diseno": "diseño",
        "tecnica": "técnica",
        "tecnico": "técnico",
        "tecnica": "técnica",
        "validacion": "validación",
        "basico": "básico",
        "practica": "práctica",
        "generacion": "generación",
        "conexion": "conexión",
        "rapidamente": "rápidamente",
        "satisfaccion": "satisfacción",
        "jerarquia": "jerarquía",
        "fisica": "física",
        "fisico": "físico",
        "interaccion": "interacción",
        "comunicacion": "comunicación",
        "unica": "única",
        "politicas": "políticas",
        "fotografica": "fotográfica",
        "fotografia": "fotografía",
        "introduccion": "introducción",
        "identificacion": "identificación",
        "solucion": "solución",
        "aplicacion": "aplicación",
        "Funcion": "Función",
        "funcion": "función",
        "segun": "según",
        "tambien": "también",
        "sesion": "sesión",
        "codigo": "código",
        "anos": "años",
        "Ano": "Año",
        "ano": "año",
        " guia": " guía",
        " guia.": " guía.",
        " Como ": " Cómo ",
        " Que ": " Qué ",
    }
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"):
                            continue
                        txt = run.text
                        if not txt:
                            continue
                        for src, dst in replacements.items():
                            txt = txt.replace(src, dst)
                        run.text = txt


def main() -> None:
    doc = Document(SRC)
    t = doc.tables

    # III.7 Analisis funcional
    fill_row(t[7], 1, [
        "Plataforma educativa web responsive con identidad de la I.E. San Jose, accesible desde celular, laptop o tablet. Integra IA, colaboracion, biblioteca, calendario, examenes y bienestar.",
        "Resuelve dudas con Profesor IA; genera examenes; permite chat y grupos en Aprendamos Juntos; organiza eventos; comparte materiales; orienta bienestar con Consejero IA y Nutrirecetas.",
        "El estudiante ingresa con cuenta, selecciona un modulo, conversa con la IA o con sus companeros, sube/consulta materiales, programa actividades y practica con evaluaciones.",
        "Sirve para acompanar el aprendizaje fuera del aula, reducir desorganizacion, facilitar estudio colaborativo y dar apoyo educativo accesible a estudiantes de secundaria.",
    ])

    # III.8 Usabilidad
    usability = [
        "La interfaz usa tarjetas, iconos y modulos claros: Profesor IA, Examen IA, Consejero IA, Aprendamos Juntos, Hora de Actuar, Biblioteca y Nutrirecetas. El recorrido es directo desde el panel principal.",
        "Puede aprenderse por exploracion. Cada modulo comunica su funcion con titulos breves, botones visibles y formularios simples. No requiere capacitacion tecnica previa.",
        "Los errores esperados son carga de archivos no permitidos, conexion lenta, sesiones vencidas o mensajes sin enviar. Se atienden con validaciones, avisos en pantalla y rutas seguras por usuario.",
        "El usuario puede abrir una consulta de IA, entrar al chat o registrar una actividad en pocos pasos. El objetivo es que una duda escolar se convierta rapidamente en respuesta, recurso o plan de accion.",
        "La satisfaccion esperada se basa en obtener ayuda inmediata, organizacion y compania de estudio. Las pruebas con usuarios reales quedan pendientes para medirla con evidencias.",
    ]
    for i, value in enumerate(usability):
        set_cell(t[8].rows[i].cells[1], value)

    # III.9 Morfologico
    morph = [
        "Software web sin forma fisica propia; se representa como una interfaz digital adaptable a pantallas de computadora, tablet y celular.",
        "Requiere dispositivo con navegador, internet y cuenta de usuario. Sus elementos visibles son panel lateral, tarjetas de acceso, chats, calendarios, biblioteca y formularios.",
        "Predominan fondo oscuro, acento dorado institucional, blancos para jerarquia textual y colores secundarios para diferenciar modulos.",
        "Textura visual digital: paneles oscuros, bordes suaves, sombras discretas, tarjetas y estados interactivos.",
        "Diseno moderno, escolar y funcional. Prioriza lectura, organizacion por modulos y continuidad visual con la identidad de Learn Up.",
    ]
    for i, value in enumerate(morph):
        set_cell(t[9].rows[i].cells[1], value)

    # III.10 Estructural
    fill_row(t[10], 1, [
        "Frontend Next.js, autenticacion Supabase, base de datos PostgreSQL, Storage, Realtime, modulos IA, notificaciones, LiveKit para llamadas y componentes de UI.",
        "El usuario interactua desde la interfaz; Supabase gestiona cuenta, perfiles, datos y archivos; la IA procesa solicitudes; Realtime actualiza mensajes/notificaciones; LiveKit conecta salas.",
        "Separar experiencia, datos, IA y comunicacion en tiempo real para que la plataforma sea escalable, segura y util para el estudiante.",
    ])

    # III.11 Comparativo: cumplir el formato sin mencionar competidores por instruccion del usuario.
    fill_row(t[11], 1, ["Producto del proyecto", "Learn Up", "Plataforma educativa integral con IA y colaboracion escolar."])
    fill_row(t[11], 2, ["Producto comparado", "PENDIENTE", "No se agregan comparaciones externas por indicacion del usuario."])
    fill_row(t[11], 4, [
        "Similitudes",
        "Diferencias",
        "Ventajas de nuestro producto/servicio",
    ])

    # Rubricas de preparacion.
    for table, label in [(t[12], "Aplicacion individual"), (t[13], "Aplicacion en equipo")]:
        set_cell(table.rows[1].cells[1], "16/05/2026", size=8)
        set_cell(table.rows[1].cells[4], "Preparacion y creacion inicial", size=8)
        levels = ["3", "3", "3", "4", "3", "3", "3", "4"]
        for r, level in zip(range(4, 12), levels):
            set_cell(table.rows[r].cells[5], f"Nivel {level}", size=8, bold=True)
        append_text(table.rows[2].cells[0], f"Observacion: {label} preliminar. Evidencia de reunion presencial con todo el equipo: PENDIENTE.", size=8)

    # Entrevista / Empatizar.
    set_cell(t[14].rows[1].cells[0], "PENDIENTE: subir a Google Drive los videos de entrevista disponibles en E:\\EPT\\Entrevista video PARTE 1.mp4 y PARTE 2.mp4.", size=8)
    questions = [
        "Que haces cuando estudias o haces tareas en casa y encuentras un tema dificil que no entiendes?",
        "Como te sientes cuando tienes una duda urgente fuera del horario de clases y no hay un profesor cerca?",
        "Como describirias tu preparacion para un examen importante y que recursos te faltan?",
        "Como te gustaria buscar y compartir resumenes, apuntes o libros digitales con tus companeros?",
        "Que dificultades tienes para organizar examenes, tareas y materiales escolares?",
        "Como te organizas con tus companeros para estudiar o hacer trabajos grupales a distancia?",
        "Como manejas el estres o frustracion por presion academica?",
        "Que habitos de alimentacion o descanso intentas seguir para rendir mejor?",
        "Que aplicaciones usas para estudiar y que les falta?",
        "Que funciones tendria la aplicacion perfecta para acompanarte en tu vida escolar?",
    ]
    for i, q in enumerate(questions, start=1):
        set_cell(t[15].rows[i].cells[1], q, size=8)
    fill_evidence_table(t[16], [
        (None, "PENDIENTE: insertar fotografia fechada o enlace de evidencia de entrevista real con consentimiento."),
        (None, "PENDIENTE: insertar evidencia de aplicacion de entrevista a estudiantes reales."),
        (None, "Documento local disponible: E:\\EPT\\Entrevistas.pdf. Para entrega final, adjuntar como anexo o colocar enlace de Google Drive."),
    ])
    set_cell(t[17].rows[1].cells[0], "Se uso entrevista semiestructurada basada en Design Thinking. Las preguntas buscaron comprender dificultades reales de estudio en casa, organizacion, colaboracion, bienestar y uso de herramientas digitales. La evidencia audiovisual con usuarios reales queda pendiente de enlace.", size=8)

    insights = [
        ("Dificultad para resolver dudas academicas fuera del horario de clase.", "Alta"),
        ("Uso de Google, YouTube o IA generica sin acompanamiento pedagogico adaptado.", "Alta"),
        ("Necesidad de organizar tareas, examenes, materiales y habitos.", "Alta"),
        ("Interes por compartir apuntes o recursos en espacios digitales visibles.", "Media"),
        ("Problemas para estudiar en grupo a distancia sin llamadas largas o desordenadas.", "Media"),
        ("Estres/frustracion academica y necesidad de apoyo emocional seguro.", "Media"),
        ("Deseo de recibir alertas, recordatorios y funciones que motiven a estudiar.", "Alta"),
    ]
    for r, (insight, freq) in enumerate(insights, start=1):
        fill_row(t[18], r, [insight, freq])

    fill_row(t[19], 1, [
        "Estudiante de secundaria de la I.E. San Jose",
        "Necesita",
        "acompanamiento academico, organizacion y apoyo colaborativo desde casa",
        "Porque",
        "fuera del horario escolar no siempre tiene docente disponible ni una herramienta integrada para estudiar.",
    ])
    fill_row(t[19], 3, [
        "Los estudiantes de secundaria necesitan una plataforma gratuita y accesible que combine IA educativa, organizacion y aprendizaje colaborativo, porque las herramientas actuales que usan estan dispersas y no responden completamente a su realidad escolar.",
        "",
        "",
        "",
        "",
    ])
    set_cell(t[20].rows[1].cells[0], "Como podriamos desarrollar una plataforma educativa gratuita, segura y accesible desde el celular que acompane a estudiantes de secundaria en sus dudas academicas, organizacion, bienestar y estudio colaborativo fuera del aula?", size=8)
    fill_evidence_table(t[21], [
        (None, "PENDIENTE: fotografia fechada de reunion de definicion del problema con el equipo."),
        (None, "PENDIENTE: fotografia de mapa de empatia o sintesis realizada por el equipo."),
        (None, "PENDIENTE: fotografia de pizarra/post-its con la formulacion del POV."),
    ])

    # Idear.
    fill_evidence_table(t[22], [
        (None, "PENDIENTE: fotografia fechada del equipo aplicando lluvia de ideas o tecnica creativa."),
        (None, "PENDIENTE: evidencia de debate de ideas de solucion."),
        (IMG_ROOT / "dashboard.png", "Evidencia digital del concepto elegido: panel principal con modulos educativos integrados."),
    ])
    set_cell(t[23].rows[1].cells[0], "Se aplico lluvia de ideas y seleccion por criterios: impacto educativo, factibilidad tecnica, acceso gratuito, utilidad para secundaria y seguridad de datos. De las ideas surgieron los modulos Profesor IA, Examen IA, Consejero IA, Aprendamos Juntos, Biblioteca, Calendario y Nutrirecetas.", size=8)
    fill_evidence_table(t[24], [
        (IMG_ROOT / "dashboard.png", "Primer esquema funcional del producto: tablero central con acceso a los modulos principales."),
        (IMG_ROOT / "ai_profesor.png", "Modulo Profesor IA como solucion principal para resolver dudas y trabajar con documentos."),
        (IMG_ROOT / "chat_oscar.png", "Modulo Aprendamos Juntos para colaboracion y mensajes entre estudiantes."),
    ])

    # Prototipar.
    fill_evidence_table(t[25], [(IMG_ROOT / "dashboard.png", "Boceto/prototipo visual de la pantalla de inicio de Learn Up.")])
    fill_evidence_table(t[26], [(IMG_ROOT / "ai_profesor.png", "Primer prototipo funcional del tutor IA para iniciar conversaciones y cargar archivos.")])
    fill_evidence_table(t[27], [(IMG_ROOT / "library.png", "Proceso de desarrollo del modulo Biblioteca para subir y compartir recursos educativos.")])
    fill_evidence_table(t[28], [(IMG_ROOT / "calendar.png", "Prueba inicial del calendario personal y compartido para organizar actividades escolares.")])

    materials = [
        ("01", "Computadora/laptop", "Herramienta", "Desarrollo, pruebas, documentacion y gestion del proyecto.", "1"),
        ("02", "Next.js / React / TypeScript", "Herramienta", "Construccion de interfaz web y logica del frontend.", "1"),
        ("03", "Supabase", "Herramienta", "Autenticacion, base de datos, almacenamiento, realtime y seguridad.", "1"),
        ("04", "APIs de IA", "Herramienta", "Generacion de respuestas, examenes, documentos e imagenes segun modulo.", "1"),
        ("05", "LiveKit", "Herramienta", "Comunicacion en llamadas y videollamadas educativas.", "1"),
        ("06", "Render", "Herramienta", "Despliegue de la aplicacion web.", "1"),
        ("07", "Canva / capturas del sistema", "Herramienta", "Diseno de infografias, publicidad y evidencias visuales.", "1"),
    ]
    for r, row in enumerate(materials, start=1):
        fill_row(t[29], r, list(row))

    # Evaluar con usuarios: pendiente.
    fill_evidence_table(t[30], [
        (None, "PENDIENTE: evidencia de prueba del prototipo con usuarios reales."),
        (None, "PENDIENTE: fotografia o captura de retroalimentacion recibida de estudiantes."),
        (None, "PENDIENTE: registro de consentimiento o autorizacion para uso de evidencia."),
    ])
    set_cell(t[31].rows[1].cells[0], "Tecnica prevista: prueba de usabilidad con estudiantes de secundaria. Se observara si comprenden los modulos, si logran enviar mensajes, consultar IA, organizar una tarea y encontrar recursos. Esta fase queda pendiente hasta realizar pruebas con usuarios reales.", size=8)
    fill_evidence_table(t[32], [
        (None, "PENDIENTE: evidencia fotografica de evaluacion con usuario real."),
        (None, "PENDIENTE: captura de encuesta o ficha de observacion aplicada."),
        (None, "PENDIENTE: evidencia de ajustes realizados despues de la prueba con usuarios."),
    ])

    # Lean Canvas. The template uses merged cells, so each box keeps its label.
    set_cell(t[33].rows[0].cells[0], "Problema\nEstudiantes sin apoyo inmediato fuera de clase; desorganizacion de tareas/examenes; dificultad para estudiar en grupo; uso disperso de herramientas digitales.", size=7)
    set_cell(t[33].rows[0].cells[1], "Solucion\nTutores IA, calendario inteligente, Aprendamos Juntos, biblioteca, examenes personalizados, consejero y nutrirecetas.", size=7)
    set_cell(t[33].rows[1].cells[1], "Metricas clave\nUsuarios activos, dudas resueltas, materiales compartidos, eventos creados, examenes generados y retencion semanal.", size=7)
    set_cell(t[33].rows[0].cells[2], "Propuesta unica de valor\nAcompanamiento educativo integral, gratuito y accesible desde una sola plataforma para estudiantes de secundaria.", size=7)
    set_cell(t[33].rows[0].cells[4], "Ventaja competitiva injusta\nConocimiento directo de la realidad de estudiantes de la I.E. San Jose y desarrollo propio adaptado a sus necesidades.", size=7)
    set_cell(t[33].rows[1].cells[4], "Canales\nPagina web, redes sociales, docentes, grupos escolares y recomendaciones entre estudiantes.", size=7)
    set_cell(t[33].rows[0].cells[5], "Segmento de clientes\nEstudiantes de secundaria de 12 a 17 anos; familias que buscan apoyo educativo accesible; docentes que desean herramientas de seguimiento.", size=7)
    set_cell(t[33].rows[2].cells[0], "Estructura de costes\nHosting, dominio, APIs de IA, base de datos, almacenamiento, notificaciones, mantenimiento y mejora de seguridad.", size=7)
    set_cell(t[33].rows[2].cells[3], "Flujo de ingresos\nModelo freemium: acceso basico gratuito; funciones premium futuras para mayor personalizacion, reportes y herramientas avanzadas.", size=7)

    set_cell(t[34].rows[1].cells[0], "Orden de validacion: 1) problema y segmento mediante entrevistas; 2) solucion con prueba de prototipo; 3) usabilidad con tareas guiadas; 4) viabilidad tecnica y costos; 5) interes por funciones premium. Las pruebas con usuarios reales quedan pendientes de evidencia formal.", size=8)
    gantt = [
        ("Entrevistas y analisis del problema", "Semana 1", "Semana 2", "2 sem.", "Equipo", "X", "X", "", "", "", "", ""),
        ("Definicion del reto y mapa de insights", "Semana 2", "Semana 3", "2 sem.", "Equipo", "", "X", "X", "", "", "", ""),
        ("Ideacion y seleccion de solucion", "Semana 3", "Semana 3", "1 sem.", "Equipo", "", "", "X", "", "", "", ""),
        ("Prototipo web inicial", "Semana 4", "Semana 5", "2 sem.", "Desarrollo", "", "", "", "X", "X", "", ""),
        ("Pruebas tecnicas internas", "Semana 5", "Semana 6", "2 sem.", "Desarrollo", "", "", "", "", "X", "X", ""),
        ("Pruebas con usuarios reales", "Pendiente", "Pendiente", "Pendiente", "Equipo", "", "", "", "", "", "", "P"),
        ("Ajustes del PMV y evidencias", "Semana 6", "Semana 7", "2 sem.", "Equipo", "", "", "", "", "", "X", "X"),
    ]
    for r, row in enumerate(gantt, start=2):
        fill_row(t[35], r, list(row), size=7)

    # Ejecucion.
    set_cell(t[36].rows[1].cells[0], "PENDIENTE: grabar/subir entrevista de problemas con usuarios reales a Google Drive.", size=8)
    set_cell(t[37].rows[1].cells[0], "PENDIENTE: grabar/subir entrevista de soluciones con usuarios reales a Google Drive.", size=8)
    validations = [
        ("01", "Acceso al dashboard", "Funciona como punto de entrada a los modulos principales."),
        ("02", "Profesor IA", "Permite iniciar conversacion y preparar soporte para archivos/documentos."),
        ("03", "Aprendamos Juntos", "Permite visualizar contactos, mensajes y llamadas; requiere seguir puliendo nombres/perfiles."),
        ("04", "Calendario", "Permite organizar actividades personales y compartidas."),
        ("05", "Biblioteca", "Permite preparar subida/gestion de materiales educativos."),
        ("06", "Seguridad de sesiones", "Se implemento control de cierre local, otros dispositivos y todos."),
        ("07", "Notificaciones", "Se agrego base de notificaciones realtime y push; falta prueba completa en dispositivos reales."),
    ]
    for r, row in enumerate(validations, start=1):
        fill_row(t[38], r, list(row))
    changes = [
        ("01", "Perfiles en Aprendamos Juntos", "Unificar datos reales desde el perfil personalizado."),
        ("02", "Cierre de sesion", "Evitar que cerrar un navegador cierre automaticamente otros dispositivos."),
        ("03", "LiveKit/llamadas", "Mostrar nombre/avatar y no codigos UUID en participantes."),
        ("04", "Notificaciones", "Mejorar inmediatez y metadatos para mensajes y eventos."),
        ("05", "Uploads", "Ajustar rutas por usuario para cumplir politicas de seguridad de Storage."),
        ("06", "IA", "Separar agentes especializados por Profesor, Consejero, Examenes y Nutrirecetas."),
        ("07", "UI", "Mantener estetica oscura institucional con acentos dorados y modulos diferenciados."),
    ]
    for r, row in enumerate(changes, start=1):
        fill_row(t[39], r, list(row))

    fill_evidence_table(t[40], [(IMG_ROOT / "ai_profesor.png", "Prototipo inicial: Profesor IA listo para iniciar conversacion y recibir archivos.")])
    fill_evidence_table(t[41], [(IMG_ROOT / "dashboard.png", "PMV: panel principal funcional con modulos de IA, colaboracion, calendario y biblioteca.")])

    set_cell(t[42].rows[1].cells[0], "PENDIENTE: publicar video publicitario en red social y colocar enlace.", size=8)
    set_cell(t[43].rows[1].cells[0], "PENDIENTE: segundo enlace de video o publicacion complementaria.", size=8)
    add_image(t[44].rows[1].cells[0], EPT_ROOT / "Infografia.png", width=2.2)
    set_cell(t[44].rows[1].cells[1], "Infografia local preparada en E:\\EPT\\Infografia.png. PENDIENTE: colocar enlace de publicacion.", size=8)
    fill_row(t[44], 2, ["PENDIENTE: afiche de difusion en redes.", "PENDIENTE"])
    fill_row(t[44], 3, ["PENDIENTE: historia/publicacion corta.", "PENDIENTE"])

    fill_evidence_table(t[45], [
        (None, "PENDIENTE: capturas de mensajes personalizados a primeros usuarios reales."),
        (IMG_ROOT / "calendar.png", "Metodo previsto: recordatorios y calendario compartido para seguimiento educativo."),
        (IMG_ROOT / "chat_oscar.png", "Metodo previsto: comunicacion por Aprendamos Juntos para seguimiento y apoyo."),
    ])
    fill_evidence_table(t[46], [
        (IMG_ROOT / "library.png", "Fidelizacion prevista: biblioteca con recursos utiles para que el estudiante vuelva a consultar materiales."),
        (IMG_ROOT / "ai_practice_generated.png", "Fidelizacion prevista: examenes personalizados para practica continua."),
        (None, "PENDIENTE: evidencia de usuarios recurrentes o testimonios reales."),
    ])

    fill_evidence_table(t[47], [
        (IMG_ROOT / "dashboard.png", "Estrategia de escala: ampliar modulos desde un panel unico y mantener acceso web."),
        (IMG_ROOT / "chat_oscar.png", "Estrategia de escala: fortalecer Aprendamos Juntos con grupos, llamadas y colaboracion."),
        (IMG_ROOT / "library.png", "Estrategia de escala: aumentar recursos compartidos en Biblioteca y validacion docente."),
    ])
    fill_evidence_table(t[48], [
        (IMG_ROOT / "profile.png", "Mejora: perfiles personalizados y datos academicos del estudiante."),
        (IMG_ROOT / "ai_practice_graded_result.png", "Mejora: correccion y retroalimentacion de examenes generados por IA."),
        (IMG_ROOT / "album.png", "Mejora: album/evidencias para registrar avances y recuerdos del proyecto."),
    ])

    challenges = [
        ("Unificar funciones en una sola plataforma sin saturar al usuario.", "Se organizaron modulos claros desde el dashboard.", "La experiencia debe priorizar rutas simples y lenguaje escolar."),
        ("Proteger sesiones, perfiles y archivos.", "Se endurecieron rutas por usuario, roles y cierre de sesion.", "La seguridad debe integrarse desde el diseno, no al final."),
        ("Lograr comunicacion en tiempo real.", "Se trabajo con Realtime, notificaciones y LiveKit.", "La rapidez percibida depende de backend y de buena interfaz."),
        ("Hacer IA util y no generica.", "Se separaron agentes por area: profesor, consejero, examenes y nutricion.", "Cada agente necesita instrucciones, limites y herramientas propias."),
        ("Recolectar evidencia con usuarios reales.", "Queda como actividad pendiente formal.", "Las pruebas reales son necesarias para validar impacto y no solo funcionamiento tecnico."),
    ]
    for r, row in enumerate(challenges, start=1):
        fill_row(t[49], r, list(row))

    lessons = [
        ("El problema educativo no es solo resolver tareas; tambien incluye organizacion, bienestar y acompanamiento.", "Realizar pruebas con estudiantes y ajustar funciones segun observaciones reales."),
        ("La IA necesita especializacion para ser confiable en educacion.", "Mejorar Profesor IA con documentos, citas y guias tipo cuaderno de estudio."),
        ("El chat colaborativo debe mostrar identidad real y ser inmediato.", "Pulir Aprendamos Juntos con perfiles, grupos, llamadas y notificaciones instantaneas."),
        ("La seguridad afecta directamente la experiencia.", "Seguir probando sesiones, permisos y storage en diferentes dispositivos."),
        ("El portafolio debe separar evidencia real de avances tecnicos.", "Completar pendientes con fotos fechadas, videos, enlaces y autorizaciones."),
    ]
    for r, row in enumerate(lessons, start=1):
        fill_row(t[50], r, list(row))

    fill_evidence_table(t[51], [
        (None, "PENDIENTE: fotografia fechada de todo el equipo reflexionando sobre el proceso."),
        (None, "PENDIENTE: acta o captura de reunion final del equipo."),
        (None, "PENDIENTE: evidencia de exposicion o revision con docente asesor."),
    ])

    # Add a minimal note in the index page-number cells only where still empty.
    for row in t[0].rows[1:]:
        if len(row.cells) >= 2 and not row.cells[1].text.strip():
            set_cell(row.cells[1], "Actualizar al finalizar", size=8)

    polish_spanish_text(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
