from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0xC8, 0x96, 0x1E)  # Gold

def add_title(text, size=28):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0xC8, 0x96, 0x1E)
    r.font.name = 'Calibri'

def add_subtitle(text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0xC8, 0x96, 0x1E)

def add_h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_quote(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    pf = p.paragraph_format
    pf.left_indent = Cm(1.5)
    pf.right_indent = Cm(1.5)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1A, 0x6B, 0x3C)

def add_body(text):
    p = doc.add_paragraph(text)

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml import OxmlElement
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1A1A2E')
        shading.set(qn('w:val'), 'clear')
        c.paragraphs[0].runs[0].font.name = 'Calibri'
        c._element.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph('')

# ===== PORTADA =====
for _ in range(6):
    doc.add_paragraph('')
add_title('PORTAFOLIO', 36)
add_title('CATEGORÍA A', 24)
doc.add_paragraph('')
add_title('CREA Y EMPRENDE 2026', 20)
doc.add_paragraph('')
add_subtitle('Proyecto: LEARN UP — Plataforma Educativa IA 24/7', 16)
add_subtitle('Equipo: Profetas del Código — Visionarios de la Tecnología', 13)
doc.add_paragraph('')
add_subtitle('Institución Educativa: (Completar)', 11)
add_subtitle('Docente Asesor: (Completar)', 11)
add_subtitle('UGEL / DRE: (Completar)', 11)
doc.add_page_break()

# ===== ÍNDICE =====
add_h1('ÍNDICE')
items = [
    '1. Nuestro Equipo Emprendedor',
    '2. Situación Problemática',
    '3. Desafío Inicial',
    '4. Fase Empatizar — Entrevistas (10)',
    '5. Fase Definir — Análisis de Información',
    '6. Desafío Final',
    '7. Fase Idear — Alternativas de Solución',
    '8. Resumen del Proyecto',
    '9. Propuesta de Valor',
    '10. Ingrediente Mágico / Innovación',
    '11. Análisis de Competidores (3)',
    '12. Fase Prototipar — PMV',
    '13. Fase Evaluar — Malla Receptora',
    '14. Lean Canvas — Modelo de Negocio',
    '15. Video Promocional',
    '16. Evidencias Fotográficas',
]
for it in items:
    add_body(it)
doc.add_page_break()

# ===== 1. EQUIPO =====
add_h1('1. NUESTRO EQUIPO EMPRENDEDOR')
add_h2('Nombre del equipo: Profetas del Código')
add_h2('¿Qué representa?: Visionarios de la tecnología')
add_body('Somos estudiantes que creen en el poder del código y la inteligencia artificial para transformar la educación. Nuestro nombre refleja la visión de anticiparnos al futuro del aprendizaje, creando herramientas tecnológicas que democraticen el conocimiento.')
doc.add_paragraph('')
add_h2('Integrantes y Roles')
add_table(
    ['Integrante', 'Rol', 'Funciones'],
    [
        ['Oscar Adrián Chanduvi Mechato', 'Líder & Programador', 'Dirige al equipo, programa y desarrolla la plataforma. Planifica sprints y asegura cohesión.'],
        ['Cristopher Josué Casanova Abad', 'Investigador', 'Investiga necesidades del público, analiza entrevistas, busca fuentes y fundamenta decisiones.'],
        ['Jharit Smart Pintado Castillo', 'Diseñador UI/UX', 'Diseña la interfaz, crea prototipos, define la experiencia de usuario y material gráfico.'],
        ['Jesús Timana Inocencio', 'Estratega & Analista', 'Analiza competencia, identifica oportunidades, diseña estrategias de comunicación.'],
        ['Cristhian J.P. Ramirez Huancas', 'Estratega & Analista', 'Investiga tendencias EdTech, diseña propuestas de valor y planifica escalabilidad.'],
        ['Santiago E. De la Cruz Jiménez', 'Portavoz', 'Representa al equipo, realiza sustentaciones, comunica la propuesta de valor.'],
    ]
)
add_h2('¿Por qué nos unimos?')
add_body('Somos un grupo de estudiantes apasionados por la tecnología y preocupados por la desigualdad educativa en nuestro país. Cada uno aporta habilidades complementarias: desde la programación hasta la comunicación, formando un equipo integral.')
doc.add_page_break()

# ===== 2. SITUACIÓN PROBLEMÁTICA =====
add_h1('2. SITUACIÓN PROBLEMÁTICA')
add_h2('Contexto Nacional')
add_body('En el Perú, la brecha educativa y digital afecta a millones de estudiantes:')
add_bullet('Solo el 20.5% de hogares rurales tiene acceso a Internet (INEI, 2025)')
add_bullet('Apenas el 8.2% de hogares rurales cuenta con computadora')
add_bullet('Más de 50,000 escuelas públicas afectadas por la falta de acceso digital')
add_bullet('Solo el 58.9% de hogares tiene conexión a Internet a nivel nacional')
add_h2('Contexto Local')
add_bullet('Falta de acompañamiento personalizado fuera del horario escolar')
add_bullet('Métodos de estudio limitados sin herramientas interactivas')
add_bullet('Aislamiento en el aprendizaje — estudian solos')
add_bullet('Desmotivación por falta de recursos dinámicos')
add_bullet('Dependencia de profesores particulares costosos')
add_h2('Problema Central')
add_quote('Los estudiantes de secundaria no cuentan con una herramienta accesible, personalizada y disponible 24 horas que les permita resolver dudas, organizar su aprendizaje y estudiar de forma colaborativa.')
doc.add_page_break()

# ===== 3. DESAFÍO INICIAL =====
add_h1('3. DESAFÍO INICIAL')
add_quote('¿Cómo podríamos nosotros crear una solución tecnológica accesible que brinde acompañamiento educativo personalizado las 24 horas a estudiantes de secundaria, permitiéndoles resolver dudas, organizar sus estudios y aprender de forma colaborativa, sin depender exclusivamente de un profesor presencial?')
doc.add_page_break()

# ===== 4. ENTREVISTAS =====
add_h1('4. FASE EMPATIZAR — ENTREVISTAS')
add_body('⚠️ Las 10 entrevistas serán completadas con evidencias reales por el equipo.')
add_h2('Guía de Preguntas para Estudiantes')
for i, q in enumerate([
    '¿Qué haces cuando tienes una duda académica fuera del horario escolar?',
    '¿Usas alguna herramienta digital para estudiar? ¿Cuál?',
    '¿Qué es lo que más te frustra al estudiar solo/a?',
    '¿Te gustaría tener un asistente virtual que te ayude con tus tareas?',
    '¿Estudias con amigos fuera del colegio? ¿Cómo se organizan?',
    '¿Qué tipo de contenido te ayudaría más?',
    '¿Tienes acceso a Internet y un dispositivo en casa?',
    '¿Cuánto tiempo dedicas al estudio diario fuera del colegio?',
    '¿Qué materia se te hace más difícil y por qué?',
    'Si existiera una app que te ayude a estudiar, ¿qué funciones necesitarías?',
], 1):
    add_bullet(f'{q}')

add_h2('Preguntas para Padres')
for q in ['¿Puede ayudar a su hijo/a con tareas?', '¿Ha considerado un profesor particular?', '¿Confiaría en IA para apoyar el aprendizaje?']:
    add_bullet(q)

add_h2('Preguntas para Docentes')
for q in ['¿Principales dificultades en sus estudiantes?', '¿La tecnología puede complementar la enseñanza?', '¿Qué herramienta digital recomendaría?']:
    add_bullet(q)

add_h2('Registro de Entrevistas')
rows = [[str(i), '(Completar)', t, '(Fecha)', '(Completar)'] for i, t in enumerate([
    'Estudiante','Estudiante','Estudiante','Estudiante',
    'Padre/Madre','Padre/Madre','Padre/Madre',
    'Docente','Docente','Estudiante'], 1)]
add_table(['N°', 'Entrevistado', 'Tipo', 'Fecha', 'Hallazgo'], rows)
add_body('📸 EVIDENCIAS FOTOGRÁFICAS: (El equipo adjuntará fotos fechadas aquí)')
doc.add_page_break()

# ===== 5. FASE DEFINIR =====
add_h1('5. FASE DEFINIR — ANÁLISIS')
add_h2('Hallazgos Clave')
add_bullet('La mayoría recurre a YouTube/Google pero no encuentra respuestas adaptadas al currículo peruano')
add_bullet('Los padres no pueden ayudar con tareas de secundaria')
add_bullet('Los docentes no pueden atender dudas individuales fuera del horario')
add_bullet('Los estudiantes desean una herramienta desde el celular, 24/7, para estudiar con amigos')
add_bullet('Profesores particulares cuestan S/. 30-80/hora, prohibitivo para la mayoría')
add_h2('Punto de Vista (POV)')
add_quote('Los estudiantes de secundaria necesitan apoyo educativo personalizado, colaborativo y 24/7 porque el sistema actual no atiende sus necesidades fuera del aula, y las alternativas son costosas, genéricas o en otro idioma.')
doc.add_page_break()

# ===== 6. DESAFÍO FINAL =====
add_h1('6. DESAFÍO FINAL')
add_quote('¿Cómo podríamos nosotros desarrollar una plataforma web educativa con IA que brinde tutoría personalizada 24/7, permita aprendizaje colaborativo y se adapte al contexto peruano, accesible desde cualquier dispositivo?')
doc.add_page_break()

# ===== 7. FASE IDEAR =====
add_h1('7. FASE IDEAR — ALTERNATIVAS')
add_table(['N°','Idea','Ventajas','Desventajas'], [
    ['1','App de flashcards con IA','Fácil de usar','Solo memorización'],
    ['2','Canal YouTube tutoriales','Accesible, visual','No interactivo'],
    ['3','Grupo WhatsApp con tutor','Familiar','Limitado'],
    ['4','Plataforma web IA + social ✅','Personalizada, 24/7, colaborativa','Requiere desarrollo'],
    ['5','Cuadernillo con QR','Sin Internet constante','Costoso, no interactivo'],
])
add_body('Idea Seleccionada: Opción 4 — Learn Up')
doc.add_page_break()

# ===== 8. RESUMEN =====
add_h1('8. RESUMEN DEL PROYECTO')
add_table(['¿Qué problema resuelve?','¿Para quién?'], [
    ['La falta de acompañamiento educativo personalizado fuera del horario escolar, que amplía la brecha educativa en Perú.',
     'Estudiantes de secundaria (12-17 años), familias que no pueden costear tutores, docentes que buscan herramientas complementarias.']
])
add_h2('¿Cómo alivias el problema?')
add_body('Learn Up ofrece una plataforma web gratuita con un tutor IA 24/7 que explica cualquier tema, se adapta al ritmo de cada estudiante, e integra aprendizaje colaborativo: grupos de estudio, calendarios compartidos y hábitos grupales.')
doc.add_page_break()

# ===== 9. PROPUESTA DE VALOR =====
add_h1('9. PROPUESTA DE VALOR')
add_h2('¿Qué vendes?')
add_bullet('Tutores IA especializados por materia')
add_bullet('Chat colaborativo "Aprendamos Juntos"')
add_bullet('Calendarios compartidos con eventos grupales')
add_bullet('Seguimiento de hábitos de estudio')
add_bullet('Biblioteca de recursos')
add_bullet('Notificaciones inteligentes')
add_h2('Propuesta de Valor Única')
add_quote('"Learn Up es tu compañero de estudio que nunca duerme." El único tutor IA 24/7 para estudiantes peruanos que responde preguntas, te conecta con compañeros, organiza tus estudios y se adapta a TU ritmo — todo gratis.')
doc.add_page_break()

# ===== 10. INGREDIENTE MÁGICO =====
add_h1('10. INGREDIENTE MÁGICO / INNOVACIÓN')
add_body('La fusión de IA personalizada + Aprendizaje Social Colaborativo:')
add_bullet('IA que conoce a tus amigos: envía mensajes, crea eventos, sugiere sesiones de estudio')
add_bullet('Hábitos compartidos: seguimiento grupal con motivación mutua')
add_bullet('Múltiples tutores IA especializados por materia')
add_bullet('Hecho POR estudiantes PARA estudiantes peruanos')
doc.add_page_break()

# ===== 11. COMPETIDORES =====
add_h1('11. ANÁLISIS DE COMPETIDORES')
add_h2('Competidor 1: Khan Academy')
add_table(['Aspecto','Khan Academy','Learn Up'], [
    ['Precio','Gratuito','Gratuito'],
    ['Idioma','Principalmente inglés','100% español, contexto peruano'],
    ['IA','Ejercicios adaptativos básicos','Tutores IA conversacionales'],
    ['Social','Sin funciones sociales','Grupos, calendarios, hábitos'],
])
add_h2('Competidor 2: Socratic by Google')
add_table(['Aspecto','Socratic','Learn Up'], [
    ['Función','Resuelve ejercicios por foto','Tutoría integral + colaborativo'],
    ['Seguimiento','Ninguno','Progreso, hábitos y metas'],
    ['Social','Ninguno','Chat grupal integrado'],
])
add_h2('Competidor 3: Duolingo')
add_table(['Aspecto','Duolingo','Learn Up'], [
    ['Precio','Freemium','Gratuito'],
    ['Materias','Solo idiomas','Todas las materias'],
    ['IA','Adaptativo para idiomas','IA conversacional integral'],
])
add_quote('Ningún competidor combina tutoría IA + aprendizaje social + organización en una plataforma gratuita para Perú.')
doc.add_page_break()

# ===== 12. PROTOTIPAR =====
add_h1('12. FASE PROTOTIPAR — PMV')
add_body('Learn Up es una plataforma web funcional: https://learn-up-qmgx.onrender.com')
add_bullet('✅ Registro e inicio de sesión seguro')
add_bullet('✅ Chat con múltiples tutores IA especializados')
add_bullet('✅ Sistema de amistades y solicitudes')
add_bullet('✅ Chat grupal "Aprendamos Juntos"')
add_bullet('✅ Calendarios compartidos')
add_bullet('✅ Seguimiento de hábitos de estudio')
add_bullet('✅ Biblioteca de recursos')
add_bullet('✅ Notificaciones en tiempo real')
add_bullet('✅ Diseño responsivo')
add_body('Tecnologías: Next.js 16, React, TypeScript, Supabase, APIs de IA')
add_body('📸 CAPTURAS DE PANTALLA: (El equipo adjuntará screenshots aquí)')
doc.add_page_break()

# ===== 13. MALLA =====
add_h1('13. FASE EVALUAR — MALLA RECEPTORA')
add_body('⚠️ Completar con feedback real de usuarios de prueba.')
add_table(['Cosas Interesantes ✨','Críticas Constructivas 🔧'], [['(Completar)','(Completar)'],['','']])
add_table(['Preguntas Nuevas ❓','Ideas Nuevas 💡'], [['(Completar)','(Completar)'],['','']])
doc.add_page_break()

# ===== 14. LEAN CANVAS =====
add_h1('14. LEAN CANVAS')
add_table(['Bloque','Descripción'], [
    ['Problema','1) Sin apoyo fuera del aula 2) Tutores costosos 3) Plataformas genéricas/inglés'],
    ['Segmento','Estudiantes secundaria 12-17 años en Perú'],
    ['Propuesta de Valor','Tutor IA 24/7 + aprendizaje colaborativo, gratuito'],
    ['Solución','Plataforma web con tutores IA, chat grupal, calendarios, hábitos'],
    ['Canales','TikTok, Instagram, colegios, ferias educativas'],
    ['Ingresos','Gratuito inicial. Futuro: freemium, alianzas'],
    ['Costos','Hosting ~$7/mes, API IA, dominio'],
    ['Métricas','Usuarios, sesiones, mensajes IA, retención'],
    ['Ventaja','Por estudiantes peruanos, IA+social, ya funcional'],
])
doc.add_page_break()

# ===== 15. VIDEO =====
add_h1('15. VIDEO PROMOCIONAL (30s)')
add_body('⚠️ ESPACIO RESERVADO — El equipo grabará y adjuntará el video.')
add_h2('Guion Sugerido')
add_bullet('[0-5s] Estudiante frustrado de noche sin resolver un ejercicio')
add_bullet('[5-10s] Abre Learn Up, la IA responde al instante')
add_bullet('[10-15s] Se conecta con amigos en "Aprendamos Juntos"')
add_bullet('[15-20s] Marca hábitos completados, sonríe')
add_bullet('[20-25s] Logo Learn Up + Profetas del Código')
add_bullet('[25-30s] "Learn Up: Tu compañero de estudio que nunca duerme."')
doc.add_page_break()

# ===== 16. EVIDENCIAS =====
add_h1('16. EVIDENCIAS FOTOGRÁFICAS')
add_body('⚠️ ESPACIO RESERVADO — Adjuntar fotos FECHADAS de:')
add_bullet('Reuniones de equipo')
add_bullet('Realización de entrevistas')
add_bullet('Sesiones de diseño y desarrollo')
add_bullet('Pruebas con usuarios')
add_bullet('Uso real de la plataforma')
add_body('IMPORTANTE: Todas las fotos deben tener la fecha visible.')
doc.add_page_break()

# ===== PERSONAJE =====
add_h1('PERSONAJE EMPRENDEDOR LOCAL')
add_body('Nombre: (Completar)')
add_body('Emprendimiento: (Completar)')
add_body('¿Por qué nos inspira?: (Completar)')
doc.add_paragraph('')
add_subtitle('Documento elaborado por Profetas del Código — Crea y Emprende 2026', 10)

# SAVE
out = r'C:\Users\profe\CyE\Portafolio_CyE_2026_LearnUp.docx'
doc.save(out)
print(f'DOCX guardado en: {out}')
